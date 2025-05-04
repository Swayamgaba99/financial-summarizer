import os
import logging
import threading
import time
from datetime import datetime
from typing import Dict, List
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openai
from langchain_openai import OpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_community.callbacks import get_openai_callback
from llama_index.core import VectorStoreIndex

logger = logging.getLogger(__name__)

class SummaryGenerator:
    """
    Generates financial summaries from documents using RAG approach.
    Creates both 2-page and 1-page summary documents.
    """

    def __init__(self, index: VectorStoreIndex, output_dir: str, openai_api_key: str):
        """
        Initialize the summary generator.

        Args:
            index: Vector index for retrieval
            output_dir: Directory to save output documents
            openai_api_key: OpenAI API key for LLM access
        """
        logger.info("Initializing SummaryGenerator")
        self.index = index
        self.output_dir = output_dir

        # Configure OpenAI
        openai.api_key = openai_api_key
        self.llm = OpenAI(openai_api_key=openai_api_key, temperature=0.1)

        # Load summary prompts
        self.summary_prompts = {
            "business_overview": {
                "prompt": """Provide business overview:

                **Required Elements:**
                1. Company Description:
                  - Founding year and headquarters
                  - location
                  - Core business activities

                2. Product Portfolio:
                  - Major hardware products
                  - Software offering
                  - Service/subscription models

                3. Operations:
                  - Manufacturing approach
                  - Supply chain strategy & Key partnerships

                4. Human Capital:
                  - Employee count (FTE)
                  - Key executive team

                5. Financial Snapshot:
                  - Current net revenue

                **Rules:**
                - Include specific numbers and percentages
                - For each 5 part give output in one single bullet point
                - Maintain professional tone

                Context: {context_str}""",
                "query": "Provide comprehensive business overview with key metrics",
                "word_limit": 200
            },
            "business_segment_overview": {
              "prompt": """Provide Business Segment Analysis.

            **Required:**
            - Top 3 revenue-generating segments
            - For each segment, include:
              - % of total revenue
              - YoY growth %
              - Profit margin %

            **Format:**
            - [Segment Name]: [X]% of revenue | Growth: [Y]% | Margin: [Z]%

            **Instructions:**
            - List in descending order of revenue contribution
            - Quantify each metric precisely
            - Give one line desc only about top 2

            Context: {context_str}""",
              "query": "Extract revenue, growth, and margin metrics by business segment, highlighting top contributors.",
              "word_limit": 100
            },
            "geographical_breakdown": {
                "prompt": """Provide Geographical analysis:

                Show only if data available:
                1. Revenue Distribution:
                  - Americas: [X]%
                  - EMEA: [Y]%
                  - APAC: [Z]%

                2. Top 3 countries by revenue with percentages

                3. Give YoY change and highlight growth market. (1 bullet point)

                **Rules:**
                - Use specific percentages


                Context: {context_str}""",
                "query": "Extract geographical revenue breakdown with growth metrics",
                "word_limit": 75
            },
            "year_over_year_analysis": {
                "prompt": """Provide YoY performance analysis in a single concise paragraph:

                    **Rules:**
                    - Use exact percentages and numbers
                    - No generic statements
                    - Focus only on measurable YoY changes with reasons

                    Context: {context_str}""",
                    "query": "Extract specific year-over-year financial metrics and most significant performance driver",
                    "word_limit": 50
            },
            "swot_analysis": {
                "prompt": """Provide SWOT analysis:

                **Foramt**
                - Strength & Weakness [one concise paragraph with below points]:
                - [Specific competitive advantage with supporting metric]
                - [Core capability with market relevance]
                - [Specific operational/financial challenge with metric]
                - [Strategic vulnerability with market context]

                - Opportunity & Threat [one concise paragraph with below points]:
                - [Concrete market opportunity with size/growth potential]
                - [Strategic initiative with potential impact]
                - [Any competitive or regulatory threat]

                **Rules:**
                - Include quantitative metrics where possible
                - Focus on forward-looking insights

                Context: {context_str}""",
                "query": "Extract material strengths, weaknesses, opportunities and threats with specific examples",
                "word_limit": 125
            },
            "credit_rating_analysis": {
                "prompt": """Provide Credit rating analysis:

              **Required Elements:**
              - Current credit ratings from major agencies (Moody's, S&P, Fitch)
              - Recent rating changes or outlook updates
              - Key debt metrics (Debt/EBITDA, interest coverage ratio)
              - Debt maturity schedule highlights
              - Liquidity position assessment

              **Format:**
              Write a concise paragraph summarizing the company's credit profile. Include specific ratings, credit metrics, and any significant changes to the debt structure. Add a brief statement on the outlook for future rating changes based on financial projections.

              **Instructions:**
              - Keep analysis under 150 words
              - Include specific ratings and debt figures
              - Focus on most recent rating agency actions
              - Note any concerns raised by rating agencies
                Context: {context_str}""",
              "query": "Extract credit rating information, debt metrics, and rating agency opinions",
              "word_limit": 150
            }
        }

    def generate_section_summary(self, section_name: str) -> str:
        """
        Generate summary for a specific section using RAG.

        Args:
            section_name: Name of the section to generate summary for

        Returns:
            str: Generated summary text
        """
        if section_name not in self.summary_prompts:
            logger.error(f"Unknown section name: {section_name}")
            return ""

        prompt_data = self.summary_prompts[section_name]
        query = prompt_data["query"]
        word_limit = prompt_data["word_limit"]

        logger.info(f"Generating '{section_name}' summary with {word_limit} word limit")

        try:
            # Create query engine from index
            query_engine = self.index.as_query_engine()

            # Execute query to get relevant context
            start_time = time.time()
            response = query_engine.query(query)
            query_time = time.time() - start_time

            logger.info(f"Retrieved context for '{section_name}' in {query_time:.2f} seconds")

            # Create LangChain prompt
            template = prompt_data["prompt"]
            prompt = PromptTemplate(
                template=template,
                input_variables=["context_str"]
            )

            # Create chain
            chain = LLMChain(llm=self.llm, prompt=prompt)

            # Generate summary with context
            with get_openai_callback() as cb:
                summary = chain.run(context_str=response.response)

            # Log token usage
            logger.info(f"Generated '{section_name}' summary using {cb.total_tokens} tokens")

            return summary.strip()

        except Exception as e:
            logger.error(f"Failed to generate '{section_name}' summary: {str(e)}")
            return f"Error generating {section_name} summary."

    def generate_two_page_summary(self) -> Dict[str, str]:
        """
        Generate comprehensive two-page summary with all sections.

        Returns:
            Dict[str, str]: Dictionary of section summaries
        """
        logger.info("Generating two-page summary")

        summary_data = {}

        # Process each section in parallel using threads
        threads = []
        results = {}

        def process_section(section_name):
            results[section_name] = self.generate_section_summary(section_name)

        start_time = time.time()

        # Create threads for each section
        for section_name in self.summary_prompts.keys():
            thread = threading.Thread(
                target=process_section,
                args=(section_name,),
                name=f"Section-{section_name}"
            )
            threads.append(thread)
            thread.start()

        # Wait for all threads to complete
        for thread in threads:
            thread.join()

        # Collect results in proper order
        for section_name in self.summary_prompts.keys():
            formatted_name = section_name.replace('_', ' ').title()
            summary_data[formatted_name] = results[section_name]

        total_time = time.time() - start_time
        logger.info(f"Generated two-page summary in {total_time:.2f} seconds")

        # Create document
        self._create_docx_document(summary_data, "two_page_summary.docx")

        return summary_data

    def generate_one_page_summary(self, two_page_summary: Dict[str, str]) -> str:
        """
        Generate condensed one-page summary from two-page summary.

        Args:
            two_page_summary: The detailed two-page summary

        Returns:
            str: Condensed one-page summary
        """
        logger.info("Generating one-page summary")

        # Combine all sections into a context for the model
        combined_context = "\n\n".join([
            f"{section_name}:\n{content}"
            for section_name, content in two_page_summary.items()
        ])

        # Create condensed summary prompt
        condensed_prompt = PromptTemplate(
            template="""
            You are a senior financial analyst creating a highly condensed summary of a financial report.
            Using the detailed report sections below, create a executive summary that captures:

            1. The company's primary business and financial position
            2. Key performance highlights across segments
            3. The most significant strengths, weaknesses, and future outlook
            4. Critical financial metrics executives need to know

            Detailed Report:
            {detailed_report}
            """,
            input_variables=["detailed_report"]
        )

        try:
            # Create and run chain
            chain = LLMChain(llm=self.llm, prompt=condensed_prompt)

            start_time = time.time()
            with get_openai_callback() as cb:
                one_page_summary = chain.run(detailed_report=combined_context)

            generation_time = time.time() - start_time

            logger.info(f"Generated one-page summary in {generation_time:.2f} seconds using {cb.total_tokens} tokens")

            # Create document
            self._create_docx_document(one_page_summary, "one_page_summary.docx")

            return one_page_summary

        except Exception as e:
            logger.error(f"Failed to generate one-page summary: {str(e)}")
            return "Error generating one-page summary."

    def _create_docx_document(self, content, filename):
        """
        Create a well-formatted Word document from the summary content.

        Args:
            content: Summary content (dict for 2-page, string for 1-page)
            filename: Output filename
        """
        logger.info(f"Creating Word document: {filename}")

        doc = DocxDocument()

        # Set narrow margins: 0.5 inch = 914400 EMUs
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.4)
            section.right_margin = Inches(0.4)

        # Add a title
        title = "Financial Summary Report"
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add generated date
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
        date_run.italic = True
        date_run.font.size = Pt(8)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if isinstance(content, dict):
            # For 2-page summary (dictionary format)
            for section_title, section_content in content.items():
                if not section_content:  # Skip empty content
                    continue

                # Format section title
                heading_para = doc.add_paragraph()
                heading_run = heading_para.add_run(section_title)
                heading_run.bold = True
                heading_run.font.size = Pt(11)
                heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Add section content
                content_para = doc.add_paragraph(section_content)
                content_para.space_after = Pt(6)  # Reduced space after content


        else:
            content_para = doc.add_paragraph(content)

        # Save the document
        output_path = os.path.join(self.output_dir, filename)
        doc.save(output_path)
        logger.info(f"Saved document to {output_path}")
        if not os.path.exists(output_path):
            raise FileNotFoundError(f"Failed to create {filename}")